import pandas as pd
from docx import Document
import os

def getTableIdx(docxPath):
    f = open(docxPath, 'rb')
    document = Document(f)
    for p in document.paragraphs:
        if 'Zestawienie bloków NMT' in p.text:
            try:
                tabNumber = int(p.text.split(' ')[1].replace('.','').replace(':', ''))
            except ValueError as e:
                print(e)
                print('Nie przetworzyło raportu {}'.format(os.path.basename(docxPath)))

    f.close()
    try:
        return tabNumber
    except UnboundLocalError as e:
        print(e)
        print('Ostrzezenie! Nie odczytano raportu {}. Przypisano tabele 1'.format(os.path.basename(docxPath)))
        return 1

def readDocxTable(docxPath, river, nHeader = 1):
    document = Document(docxPath)
    tableNum = getTableIdx(docxPath)
    try:
        table = document.tables[tableNum]
        data = [[cell.text for cell in row.cells] for row in table.rows]

        l = len(data)
        if river:
            data[0].append('rzeka')
            data[0].append('idhydr')
            data[0].append('wersja')

            for i in range(1,l):
                data[i].append(river[0])
                data[i].append(river[-2])
                data[i].append(river[-1])
        # print(data)
        df = pd.DataFrame(data)
        if nHeader == 1:
            df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
        elif nHeader == 2:
            outside_col, inside_col = df.iloc[0], df.iloc[1]
            hier_index = pd.MultiIndex.from_tuples(list(zip(outside_col, inside_col)))
            df = pd.DataFrame(data,columns=hier_index).drop(df.index[[0,1]] ).reset_index(drop=True)
        elif nHeader > 2:
            print("Tabela posaida wiecej niż 2 naglowki")
            df = pd.DataFrame()
    except IndexError as e:
        print(e)
        print('Nie przetworzyło raportu {}'.format(os.path.basename(docxPath)))
        pass
    return df

def search(in_folder,file_ext, version):
    frames = []
    for root, dirs, files in os.walk(in_folder, topdown=False):
        for name in files:
            if name.endswith(file_ext):
                file = os.path.join(root, name)
                doc_dir = os.path.dirname(os.path.dirname(file))
                lok_doc = os.path.basename(doc_dir)
                if version in lok_doc:
                    print(lok_doc)
                    rzeka = lok_doc.split('_')
                    frame = readDocxTable(file,rzeka, 1)
                    frames.append(frame)
    return frames

if __name__ == '__main__':
    # document = r"C:\robo\Przeglad_LOKAL_ROBO\aMZPiMRP 1 3 14 27 Raport z wyznaczania OZP dla Baudy 20190219 v3.00.docx"
    folder = r'R:\10. Modele hydrauliczne\RW G-ZW'
    output_csv = r'C:\robo\Przeglad_LOKAL_ROBO\RW_GZW_wynikowy2022.csv'
    # getTableIdx(document)
    frames = search(folder, '.docx', '2022v1')
    data = pd.concat(frames)
    # data = readDocxTable(document)
    print(data.to_string())
    data.to_csv(output_csv)