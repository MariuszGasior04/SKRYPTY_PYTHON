#!/usr/bin/env python
#-*- coding: utf-8 -*-

#skrypt do ankiet.
import os
import arcpy
from arcpy import env
arcpy.env.overwriteOutput = True
import time
import docx
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.dml.color import ColorFormat

arcpy.env.overwriteOutput = True
arcpy.env.workspace = ws = ur'P:\Projekty_2017\18021-00_SZCW\07.GIS\06_Ankiety\Baza_do_ankiet\Baza_do_ankiet.gdb'
aJCWP = 'JCWP_rzeczne'

kol = []

for field in arcpy.ListFields(aJCWP):
        kol.append(str(field.name))

def add_titlerun(title_run, paragraph):

    run = paragraph.add_run(title_run)
    font = run.font
    font.name = 'Calibri'
    font.bold = True
    font.size = Pt(16)
    font.color.rgb = RGBColor(0x21, 0x58, 0x68)
    return

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

with arcpy.da.UpdateCursor(aJCWP, kol) as cur:
    for row in cur:
        if row[kol.index('KARTA')] == 2 and row[kol.index('WYZN_WST')] != 'NAT':
            print(row[kol.index('MS_KOD')])

            if row[kol.index('WYZN_WST')]==u'SZCW':
                dok = docx.Document('_Kwestionariusz_rzeki_SZCW_szablon.docx')
            elif row[kol.index('WYZN_WST')]==u'SCW':
                dok = docx.Document('_Kwestionariusz_rzeki_SCW_szablon.docx')

            """Tytuł"""
            paragraph = dok.paragraphs
            paragraph[1].clear()
            add_titlerun(row[kol.index('Nazwa_JCWP')],paragraph[1])
            paragraph[2].clear()
            add_titlerun(row[kol.index('MS_KOD')],paragraph[2])

            """pytanie 1A"""
            if row[kol.index('WPTR_P')]+row[kol.index('WPTR_G')] == 0:
                for i in range(8,61):
                    delete_paragraph(paragraph[i])

            """pytanie 1B"""
            if row[kol.index('PPH2_P')]+row[kol.index('PPH2_G')] == 0:
                for i in range(61,110):
                    delete_paragraph(paragraph[i])

            """pytanie 1C"""
            if row[kol.index('PPH3_P')]+row[kol.index('PPH3_G')] == 0:
                for i in range(110,154):
                    delete_paragraph(paragraph[i])

            """pytanie 1D"""
            if row[kol.index('PPH4_P')]+row[kol.index('PPH4_G')] == 0:
                for i in range(154,181):
                    delete_paragraph(paragraph[i])

            """pytanie 1E"""
            if row[kol.index('PPH6_P')]+row[kol.index('PPH6_G')] == 0:
                for i in range(181,198):
                    delete_paragraph(paragraph[i])

            """pytanie 1F"""
            if row[kol.index('KOREKT_EKS')] != 'TAK':
                for i in range(198,251):
                    delete_paragraph(paragraph[i])

            """Utworzenie ankiety"""
            if row[kol.index('WYZN_WST')]==u'SZCW':
                dok.save(os.path.join('zz_'+row[kol.index('ZZ')]+' '+str(row[kol.index('MS_KOD')])+' kwestionariusz SZCW'+'.docx'))
                print(str(row[kol.index('MS_KOD')])+'.docx')
            elif row[kol.index('WYZN_WST')]==u'SCW':
                dok.save(os.path.join('zz_'+row[kol.index('ZZ')]+' '+str(row[kol.index('MS_KOD')])+' kwestionariusz SCW'+'.docx'))
                print(str(row[kol.index('MS_KOD')])+'.docx')

del cur