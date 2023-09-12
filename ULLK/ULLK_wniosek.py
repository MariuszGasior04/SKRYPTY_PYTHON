#!/usr/bin/env python
#-*- coding: utf-8 -*-

#skrypt do tworzenia wsadu do wniosku

import os
import arcpy
import docx
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

arcpy.env.overwriteOutput = True

"""parametry programu"""

##workspace = arcpy.env.workspace = r"R:\OIIS_KR5\_PROJEKTY 2019\19004-00_Piekielko\33_ULLK\337_GIS\2_mapy\odc_5\baza_mapy_odc5.gdb"
##
##dzialki_czas = 'dzialki_ewidencyjne_zajecie_czasowe_diss'
##kol = []
##
##
##lista_robot = []
##with arcpy.da.SearchCursor('dzialki_czas', ['ID_DZ','ETYK1','ETYK2','ETYK3','ETYK4','ETYK5','ETYK6','ETYK7','ETYK8','ETYK9','ETYK10','ETYK11','ETYK12','ETYK13','ETYK14','ETYK15','ETYK16','ETYK17']) as cur:
##    for row in cur:
##        lista_robot.append(row)
##del cur
##
document = docx.Document()

paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
paragraph = paragraph.add_run(u'III.	W związku z art.9q ust. 1 pkt 6 oraz art. 9ya ust. 1 ustawy o transporcie kolejowym wnoszę o ustalenie ograniczenia w korzystaniu z nieruchomości w celu zapewnienia prawa do wejścia na teren nieruchomości dla prowadzenia inwestycji kolejowej stosownie do treści ww. artykułu (w tym wymagające przejścia przez tereny wód płynących bądź dróg publicznych), na następujących nieruchomościach (w nawiasach podano numery działek ewidencyjnych przed podziałem):')
font = paragraph.font
font.name = 'Arial'
font.bold = False
font.size = Pt(11)
paragraph = document.add_paragraph()



document.save(r'R:\OIIS_KR5\_PROJEKTY 2019\19004-00_Piekielko\33_ULLK\337_GIS\4_produkty\wniosek_wsad.docx')
